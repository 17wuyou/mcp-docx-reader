from dataclasses import dataclass
from typing import AsyncIterator
from mcp.server.fastmcp import FastMCP
from docx import Document
import os
from contextlib import asynccontextmanager

# Directory where DOCX files are stored
DOCX_DIRECTORY = os.getenv("DOCX_DIRECTORY", "./docx")

@dataclass
class AppContext:
    """Application context for lifecycle management."""
    docx_directory: str

# Initialize the MCP server (lifespan added below)
mcp = FastMCP("DOCX Reader")

@asynccontextmanager
async def app_lifespan(server: FastMCP) -> AsyncIterator[AppContext]:
    """Manage application lifecycle with type-safe context"""
    try:
        # Setup can go here
        yield AppContext(docx_directory=DOCX_DIRECTORY)
    finally:
        # Cleanup (if needed)
        pass

# Assign lifespan to server
mcp.lifespan = app_lifespan

@mcp.tool()
def read_docx(ctx, filename: str) -> str:
    """
    Reads and extracts text from a specified DOCX file.
    :param ctx: FastMCP context
    :param filename: Name of the DOCX file to read
    :return: Extracted text from the DOCX
    """
    # 获取docx目录，优先使用上下文中的目录，否则使用全局变量
    try:
        docx_directory = getattr(ctx, 'docx_directory', DOCX_DIRECTORY)
    except:
        docx_directory = DOCX_DIRECTORY
    
    docx_path = os.path.join(docx_directory, filename)

    if not os.path.exists(docx_path):
        return f"Error: File '{filename}' not found at {docx_path}."

    try:
        # Open and extract text from the DOCX
        doc = Document(docx_path)
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text.append(paragraph.text)
        
        # Add table content if present
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = ' '.join([p.text for p in cell.paragraphs]).strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text.append('\t'.join(row_text))
        
        extracted_text = "\n".join(text)
        return extracted_text if extracted_text else "No text found in the DOCX."
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"

# Run the MCP server
def main():
    mcp.run()

if __name__ == "__main__":
    main()